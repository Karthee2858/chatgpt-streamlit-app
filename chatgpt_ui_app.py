import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from weasyprint import HTML
from datetime import datetime
import os

def fetch_chat_content(url):
    try:
        res = requests.get(url)
        soup = BeautifulSoup(res.text, 'html.parser')
        messages = []

        for block in soup.select('[data-message-author-role]'):
            role = block['data-message-author-role']
            text_block = block.select_one('div.markdown') or block.select_one('div.prose')
            text = text_block.get_text("\n") if text_block else ''
            messages.append((role.capitalize(), text.strip()))

        return messages
    except Exception as e:
        return []

def save_as_word(messages, filename):
    doc = Document()
    doc.add_heading("ChatGPT Conversation", 0)
    for role, text in messages:
        doc.add_paragraph(f"{role}:", style='Heading 2')
        doc.add_paragraph(text)
    doc.save(filename)
    return filename

def save_as_pdf(messages, filename):
    html_content = "<h1>ChatGPT Conversation</h1>"
    for role, text in messages:
        html_content += f"<h2>{role}</h2><p>{text.replace('\n', '<br>')}</p>"
    HTML(string=html_content).write_pdf(filename)
    return filename

def main():
    st.set_page_config(page_title="ChatGPT Memory Saver", layout="centered")
    st.title("ChatGPT Memory Saver")
    st.write("Convert your ChatGPT conversation into PDF or Word for offline study!")

    url = st.text_input("Paste your ChatGPT share link")

    file_type = st.radio("Choose file format:", ("Word (.docx)", "PDF (.pdf)"))

    if st.button("Generate File"):
        if not url.strip():
            st.error("Please enter a valid ChatGPT link.")
            return

        with st.spinner("Fetching and converting..."):
            messages = fetch_chat_content(url)
            if not messages:
                st.error("Could not fetch messages. Check your link.")
                return

            date_str = datetime.now().strftime("%Y-%m-%d")
            base_name = f"ChatGPT_Conversation_{date_str}"
            output_file = ""

            if file_type == "Word (.docx)":
                output_file = save_as_word(messages, base_name + ".docx")
            else:
                output_file = save_as_pdf(messages, base_name + ".pdf")

            with open(output_file, "rb") as file:
                btn = st.download_button(
                    label="Download File",
                    data=file,
                    file_name=os.path.basename(output_file),
                    mime="application/octet-stream"
                )

            st.success("File generated successfully!")

if __name__ == "__main__":
    main()
